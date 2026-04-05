import os
import re
import glob
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Using pypdf for robust PDF text extraction
try:
    from pypdf import PdfReader
except ImportError:
    print("Please install pypdf: pip install pypdf")
    exit()

# --- CONFIGURATION ---
BASE_DIR = os.getcwd()
OUTPUT_DIR = os.path.join(BASE_DIR, "Generated_Applications")

def extract_text_from_pdf(pdf_path):
    """Synchronously extracts all text from a PDF file."""
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def parse_cv_data(text):
    """
    Parses contact info and highlights from the CV text using German-specific patterns.
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # Extract Email
    email = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    
    # Extract German Phone (e.g., +49 or 0151...)
    phone = re.search(r'(\+49|0)[1-9][0-9 \-\/]{7,}', text)
    
    # Extract German Address (Looks for 5-digit ZIP + City)
    address_match = re.search(r'([A-ZÄÖÜa-z\s\.\-]+?\d{1,4}[a-z]?\s*,?\s*\d{5}\s+[A-ZÄÖÜa-z\s\-]+)', text)
    
    # Basic Name Extraction (usually the first line of a CV)
    name = lines[0] if lines else "Vorname Nachname"

    return {
        "name": name,
        "email": email.group(0) if email else "ihre.mail@beispiel.de",
        "phone": phone.group(0) if phone else "Mobil: 0123 456789",
        "address": address_match.group(0) if address_match else "Musterstraße 1, 12345 Stadt",
        "raw_content": text[:2000] # Save snippet for context
    }

def create_letter(index, job_row, user_info):
    """Generates a native-level German cover letter synchronously."""
    doc = Document()
    
    # Job details from CSV
    company = str(job_row.get('Company', 'Unternehmen')).strip()
    position = str(job_row.get('Position', 'Ihre ausgeschriebene Stelle')).strip()
    city = str(job_row.get('Location', 'Deutschland')).strip()

    # --- STYLE SETUP ---
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # --- SENDER HEADER ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run(f"{user_info['name']}\n{user_info['address']}\n{user_info['phone']}\n{user_info['email']}")

    # --- RECIPIENT ---
    doc.add_paragraph(f"\n{company}\nPersonalabteilung\n{city}")

    # --- DATE ---
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    current_date = datetime.now().strftime('%d.%m.%Y')
    date_p.add_run(f"\n{city}, den {current_date}")

    # --- SUBJECT (BETREFF) ---
    subject = doc.add_paragraph(f"\nBewerbung als {position}")
    subject.runs[0].bold = True
    subject.runs[0].font.size = Pt(12)

    # --- BODY ---
    # Salutation
    doc.add_paragraph("\nSehr geehrte Damen und Herren,")

    # Introduction
    p1 = doc.add_paragraph(
        f"mit großem Interesse verfolge ich die aktuelle Entwicklung von {company} "
        f"und bewerbe mich hiermit auf die von Ihnen ausgeschriebene Position als {position}. "
        "Die beschriebenen Aufgabenfelder decken sich hervorragend mit meinem bisherigen "
        "Werdegang und meinen fachlichen Kompetenzen."
    )
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Experience Section (Dynamic phrasing)
    p2 = doc.add_paragraph(
        "In meiner bisherigen beruflichen Laufbahn konnte ich fundierte Erfahrungen "
        "sammeln, die ich nun gewinnbringend in Ihr Team einbringen möchte. "
        "Ich verfüge über eine strukturierte Arbeitsweise, ein hohes Maß an "
        "Eigeninitiative und die Fähigkeit, auch in stressigen Phasen stets den "
        "Überblick zu behalten."
    )
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Closing
    p3 = doc.add_paragraph(
        "Gerne möchte ich Sie in einem persönlichen Gespräch von meiner Motivation "
        "und meiner Eignung für diese Position überzeugen. Über eine positive Rückmeldung "
        "Ihrerseits freue ich mich sehr."
    )
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Sign-off
    doc.add_paragraph("\nMit freundlichen Grüßen\n\n")
    doc.add_paragraph(user_info['name'])

    # Save
    safe_name = re.sub(r'\W+', '', company)
    file_path = os.path.join(OUTPUT_DIR, f"{index:02d}_Bewerbung_{safe_name}.docx")
    doc.save(file_path)

def main():
    # 1. Setup Directories
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    # 2. Find Files (Synchronous search)
    cv_files = glob.glob(os.path.join(BASE_DIR, "*.pdf"))
    csv_files = glob.glob(os.path.join(BASE_DIR, "*.csv"))

    if not cv_files or not csv_files:
        print("Error: Ensure a .pdf CV and a .csv Job list are in the directory.")
        return

    # 3. Synchronous PDF Extraction
    print(f"Reading CV: {os.path.basename(cv_files[0])}...")
    cv_text = extract_text_from_pdf(cv_files[0])
    user_info = parse_cv_data(cv_text)

    # 4. Synchronous Job Processing
    print(f"Reading Jobs: {os.path.basename(csv_files[0])}...")
    jobs_df = pd.read_csv(csv_files[0])

    # Check for required columns
    required = ['Company', 'Position', 'Location']
    for col in required:
        if col not in jobs_df.columns:
            # Fallback to fuzzy matching if exact names aren't found
            for df_col in jobs_df.columns:
                if col.lower() in df_col.lower():
                    jobs_df.rename(columns={df_col: col}, inplace=True)

    print(f"Generating {len(jobs_df)} documents for {user_info['name']}...")

    # Iterate through rows sequentially
    for index, row in jobs_df.iterrows():
        try:
            create_letter(index + 1, row, user_info)
            print(f"  [+] Created: {row.get('Company', 'Job_'+str(index))}")
        except Exception as e:
            print(f"  [!] Failed {index}: {e}")

    print(f"\nSuccess! Files saved in: {OUTPUT_DIR}")

if __name__ == "__main__":
    main()